"""
Trillium RAG - Pagina Chat RAG
Layout a due colonne: chat (sinistra) + pannello fonti/bibliografia (destra).
Integra: citazioni, snippet, multi-turno, export, feedback, confidence, domande suggerite.
"""

import os
import io
import streamlit as st

from rag.query import rag_query, retrieve_relevant_docs
from citation_parser import parse_citations, match_citations_to_sources
from feedback import save_rating, generate_suggested_questions
from confidence import calculate_confidence
from auth import get_user_role
from modules.helpers import (
    get_available_providers,
    short_path_for_display,
    get_file_for_download,
    extract_doc_identifier,
)


# ============================================================
# STILI CSS
# ============================================================

def _inject_chat_styles():
    st.markdown("""
    <style>
    /* Pannello fonti */
    .source-card {
        background: #f8fafc;
        border: 1px solid #e2e8f0;
        border-radius: 8px;
        padding: 12px 14px;
        margin-bottom: 10px;
        transition: box-shadow 0.2s;
    }
    .source-card:hover {
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    }
    .source-title {
        font-weight: 600;
        font-size: 14px;
        color: #1a202c;
        margin-bottom: 4px;
    }
    .source-path {
        font-size: 11px;
        color: #718096;
        margin-bottom: 6px;
    }
    .source-snippet {
        font-size: 12px;
        color: #4a5568;
        line-height: 1.4;
        margin-bottom: 8px;
    }
    .source-cited {
        border-left: 3px solid #3182ce;
    }

    /* Score bar */
    .score-container {
        display: flex;
        align-items: center;
        gap: 8px;
        margin-bottom: 6px;
    }
    .score-bar-bg {
        flex: 1;
        height: 6px;
        background: #e2e8f0;
        border-radius: 3px;
        overflow: hidden;
    }
    .score-bar-fill {
        height: 100%;
        border-radius: 3px;
        transition: width 0.3s;
    }
    .score-label {
        font-size: 12px;
        font-weight: 600;
        min-width: 38px;
        text-align: right;
    }

    /* Confidence badge */
    .confidence-badge {
        border-radius: 8px;
        padding: 10px 14px;
        margin-bottom: 14px;
    }
    .confidence-badge .conf-score {
        font-size: 20px;
        font-weight: 700;
    }
    .confidence-badge .conf-msg {
        font-size: 12px;
        margin-top: 2px;
        opacity: 0.85;
    }

    /* Chat area */
    .timing-badge {
        display: inline-block;
        background: #edf2f7;
        border-radius: 12px;
        padding: 2px 10px;
        font-size: 11px;
        color: #718096;
    }
    </style>
    """, unsafe_allow_html=True)


# ============================================================
# PANNELLO FONTI (colonna destra)
# ============================================================

def _get_score_color(score_pct):
    """Colore per la barra score in base alla percentuale."""
    if score_pct >= 80:
        return "#38a169"  # verde
    elif score_pct >= 60:
        return "#d69e2e"  # giallo
    elif score_pct >= 40:
        return "#dd6b20"  # arancione
    else:
        return "#e53e3e"  # rosso


def _render_source_panel(docs, citations=None):
    """Renderizza il pannello fonti nella colonna destra."""
    if not docs:
        st.markdown("""
        <div style='text-align:center; padding: 40px 20px; color: #a0aec0'>
            <div style='font-size: 32px; margin-bottom: 8px'>📚</div>
            <div style='font-size: 14px'>Fai una domanda per vedere<br>i documenti utilizzati</div>
        </div>
        """, unsafe_allow_html=True)
        return

    # Raggruppa documenti per source e calcola score migliore per fonte
    source_data = {}
    for d in docs:
        src = d.get("source", "")
        if not src:
            continue
        score = d.get("score", 0.0)
        if src not in source_data:
            source_data[src] = {"best_score": score, "texts": [], "doc": d}
        if score > source_data[src]["best_score"]:
            source_data[src]["best_score"] = score
        source_data[src]["texts"].append(d.get("text", ""))

    # Ordina per score decrescente
    sorted_sources = sorted(source_data.items(), key=lambda x: x[1]["best_score"], reverse=True)

    cited_paths = set()
    if citations:
        for c in citations:
            if "matched_source" in c:
                cited_paths.add(c["matched_source"])

    # Confidence badge
    confidence = st.session_state.get("last_confidence")
    if confidence:
        score = confidence["score"]
        level = confidence["level"]
        color_map = {"alta": "#38a169", "media": "#d69e2e", "bassa": "#e53e3e"}
        color = color_map.get(level, "#718096")
        bg = f"{color}18"
        st.markdown(
            f"<div class='confidence-badge' style='background:{bg}; border-left:4px solid {color}'>"
            f"<div class='conf-score' style='color:{color}'>{score}% Confidenza</div>"
            f"<div class='conf-msg'>{confidence['message']}</div></div>",
            unsafe_allow_html=True,
        )

    st.markdown(f"**{len(sorted_sources)} fonte/i** — {len(docs)} chunk recuperati")
    st.markdown("---")

    for idx, (source, data) in enumerate(sorted_sources):
        doc_id = extract_doc_identifier(source)
        short = short_path_for_display(source, max_segments=3)
        is_cited = source in cited_paths
        raw_score = data["best_score"]

        # Converti score cosine (0..1) in percentuale
        score_pct = round(raw_score * 100, 1) if raw_score <= 1.0 else round(raw_score, 1)
        score_color = _get_score_color(score_pct)

        # Card con score bar
        cited_class = "source-cited" if is_cited else ""
        cited_badge = " <span style='color:#3182ce; font-size:11px'>● citato</span>" if is_cited else ""

        # Snippet
        first_text = data["texts"][0] if data["texts"] else ""
        snippet = first_text[:180].strip().replace("\n", " ")
        if len(first_text) > 180:
            snippet += "..."

        card_html = f"""
        <div class='source-card {cited_class}'>
            <div class='source-title'>{doc_id}{cited_badge}</div>
            <div class='source-path'>{short}</div>
            <div class='score-container'>
                <div class='score-bar-bg'>
                    <div class='score-bar-fill' style='width:{score_pct}%; background:{score_color}'></div>
                </div>
                <div class='score-label' style='color:{score_color}'>{score_pct}%</div>
            </div>
            <div class='source-snippet'>{snippet}</div>
        </div>
        """
        st.markdown(card_html, unsafe_allow_html=True)

        # Download button
        is_url = source.strip().lower().startswith(("http://", "https://"))
        if is_url:
            st.markdown(f"[🔗 Apri link]({source})")
        else:
            file_info = get_file_for_download(source)
            if file_info:
                file_data, name, mime = file_info
                st.download_button(
                    f"📥 {name[:40]}",
                    data=file_data,
                    file_name=name,
                    mime=mime,
                    key=f"src_{idx}_{hash(source) % 10**8}",
                    use_container_width=True,
                )

        # Immagine anteprima
        if _is_image_file(source) and os.path.isfile(source):
            try:
                st.image(source, width=200, caption=os.path.basename(source))
            except Exception:
                pass


def _is_image_file(path: str) -> bool:
    ext = os.path.splitext(path)[1].lower()
    return ext in (".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".heic", ".heif")


# ============================================================
# CHAT HISTORY
# ============================================================

def _render_chat_history():
    if not st.session_state.chat_history:
        st.markdown("""
        <div style='text-align:center; padding: 60px 20px; color: #a0aec0'>
            <div style='font-size: 48px; margin-bottom: 12px'>🔍</div>
            <div style='font-size: 16px; font-weight: 500'>Trillium RAG Assistant</div>
            <div style='font-size: 13px; margin-top: 4px'>Fai una domanda sui tuoi documenti tecnici</div>
        </div>
        """, unsafe_allow_html=True)
        return

    for msg_idx, msg in enumerate(st.session_state.chat_history):
        if msg["role"] == "user":
            with st.chat_message("user"):
                st.write(msg["content"])
        else:
            with st.chat_message("assistant"):
                st.write(msg["content"])

                # Pulsanti feedback sotto ogni risposta dell'assistente
                fb_col1, fb_col2, fb_col3 = st.columns([1, 1, 6])
                fb_key = f"fb_{msg_idx}"
                rated_key = f"rated_{msg_idx}"

                if not st.session_state.get(rated_key):
                    with fb_col1:
                        if st.button("👍", key=f"{fb_key}_pos", help="Risposta utile"):
                            q = ""
                            for m in st.session_state.chat_history[:msg_idx]:
                                if m["role"] == "user":
                                    q = m["content"]
                            save_rating(q, msg["content"], "positive")
                            st.session_state[rated_key] = "positive"
                            st.rerun()
                    with fb_col2:
                        if st.button("👎", key=f"{fb_key}_neg", help="Risposta non utile"):
                            q = ""
                            for m in st.session_state.chat_history[:msg_idx]:
                                if m["role"] == "user":
                                    q = m["content"]
                            save_rating(q, msg["content"], "negative")
                            st.session_state[rated_key] = "negative"
                            st.rerun()
                else:
                    rating = st.session_state[rated_key]
                    with fb_col1:
                        st.caption("👍 Grazie!" if rating == "positive" else "👎 Grazie")


# ============================================================
# EXPORT WORD
# ============================================================

def _export_chat_as_word():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()
    title = doc.add_heading("Trillium RAG — Conversazione", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for msg in st.session_state.chat_history:
        role = "Domanda" if msg["role"] == "user" else "Risposta"
        p = doc.add_paragraph()
        run = p.add_run(f"{role}:")
        run.bold = True
        run.font.size = Pt(11)
        if msg["role"] == "user":
            run.font.color.rgb = RGBColor(0x34, 0x40, 0x54)
        else:
            run.font.color.rgb = RGBColor(0x10, 0x5C, 0x3E)
        p_content = doc.add_paragraph(msg.get("content", ""))
        p_content.style.font.size = Pt(10)
        doc.add_paragraph("—" * 40)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# PAGINA PRINCIPALE
# ============================================================

def render():
    _inject_chat_styles()

    col_chat, col_sources = st.columns([7, 3])

    with col_sources:
        st.markdown("### 📋 Fonti e Rilevanza")
        st.caption("ℹ️ Qui vengono mostrate le fonti utilizzate per la risposta. "
                   "Ogni fonte ha un punteggio di rilevanza (%). "
                   "Usa gli Smart Filters per affinare la ricerca. "
                   "Esempio: seleziona famiglia 'BB1' + tipo 'parts_list' per cercare solo nelle parts list delle BB1.")

        # Smart Filters per ricerca avanzata
        with st.expander("🎯 Smart Filters", expanded=False):
            st.caption("Filtra i documenti prima della ricerca semantica")

            # Filtro famiglia pompa
            pump_family = st.selectbox(
                "Famiglia Pompa",
                ["Tutte", "OH1", "OH2", "OH3", "OH4", "OH5",
                 "BB1", "BB2", "BB3", "BB4", "BB5",
                 "VS1", "VS2", "VS3", "VS4", "VS5", "VS6", "VS7"],
                index=0,
                key="filter_pump_family",
                help="Filtra per famiglia pompa API 610",
            )

            # Filtro tipo documento (arricchito)
            doc_type = st.selectbox(
                "Tipo Documento",
                ["Tutti", "parts_list", "general_arrangement", "cross_section",
                 "datasheet", "drawing", "specification", "calculation",
                 "procedure", "standard", "other"],
                format_func=lambda x: {
                    "Tutti": "📄 Tutti i documenti",
                    "parts_list": "📋 Parts List / BOM",
                    "general_arrangement": "📐 General Arrangement",
                    "cross_section": "🔍 Cross Section",
                    "datasheet": "📊 Datasheet",
                    "drawing": "✏️ Disegno tecnico",
                    "specification": "📝 Specifica",
                    "calculation": "🧮 Calcolo (Mod.)",
                    "procedure": "📘 Procedura (SOP)",
                    "standard": "📏 Normativa/Standard",
                    "other": "📁 Altro",
                }.get(x, x),
                index=0,
                key="filter_doc_type",
            )

            # Filtro peso
            has_weight = st.checkbox(
                "🏋️ Solo documenti con dati peso",
                value=False,
                key="filter_has_weight",
                help="Mostra solo documenti che contengono informazioni sul peso"
            )

            # Filtro flange rating
            flange_rating = st.selectbox(
                "Flange Rating",
                ["Qualsiasi", 150, 300, 600, 900, 1500, 2500],
                index=0,
                key="filter_flange_rating",
                help="Class ASME B16.5",
            )

            # Filtro materiale (testo libero)
            material_filter = st.text_input(
                "🔬 Materiale",
                placeholder="es. CF8M, Duplex, WCB...",
                key="filter_material",
                help="Cerca nei materiali estratti dai documenti",
            )

            st.divider()

            # Filtri legacy (SOP range)
            filter_types = st.multiselect(
                "Filtro legacy (SOP/Mod)",
                ["sop", "mod", "standard", "other"],
                format_func=lambda x: {"sop": "SOP", "mod": "Mod (Excel)", "standard": "Normative", "other": "Altro"}.get(x, x),
                default=[],
                key="filter_doc_types",
                help="Filtro basato su nome file (compatibilità)",
            )
            use_sop_range = st.checkbox("Filtra per range SOP", key="use_sop_range")
            sop_min, sop_max = 100, 999
            if use_sop_range:
                sop_min, sop_max = st.slider("Range SOP", 100, 999, (400, 600), key="sop_range")

            # Costruisci filtri combinati
            smart_filters = {}
            if pump_family != "Tutte":
                smart_filters["pump_family"] = pump_family
            if doc_type != "Tutti":
                smart_filters["doc_type"] = doc_type
            if has_weight:
                smart_filters["has_weight"] = True
            if flange_rating != "Qualsiasi":
                smart_filters["flange_rating"] = flange_rating
            if material_filter.strip():
                smart_filters["material"] = material_filter.strip()

            st.session_state.search_filters = {
                "doc_types": filter_types,
                "sop_min": sop_min if use_sop_range else None,
                "sop_max": sop_max if use_sop_range else None,
                **smart_filters,  # Aggiungi smart filters
            }

            # Badge filtri attivi
            active_count = len(smart_filters)
            if active_count > 0:
                st.success(f"✅ {active_count} filtro/i attivo/i")

        last_docs = st.session_state.get("last_retrieved_docs", [])
        last_citations = st.session_state.get("last_citations", [])
        _render_source_panel(last_docs, last_citations)

        # Dettagli tecnici
        if last_docs:
            timing = st.session_state.get("last_timing", {})
            with st.expander("⚙️ Dettagli tecnici"):
                st.caption(f"Chunk recuperati: {len(last_docs)}")
                st.caption(f"Fonti uniche: {len(set(d.get('source', '') for d in last_docs if d.get('source')))}")
                if last_citations:
                    st.caption(f"Citazioni estratte: {len(last_citations)}")
                if timing:
                    if timing.get("cache_hit"):
                        st.caption("⚡ Risposta dalla cache")
                    for k, v in timing.items():
                        if k not in ("cache_hit",) and isinstance(v, (int, float)):
                            st.caption(f"{k}: {v}s")

    with col_chat:
        # Selezione provider
        providers = get_available_providers()
        if not providers:
            st.error("Nessun provider LLM configurato.")
            st.stop()

        provider_options = [f"{name} ({model})" for name, model, _ in providers]
        selected_idx = 0
        for i, (_, _, code) in enumerate(providers):
            if code == st.session_state.get("selected_provider"):
                selected_idx = i
                break

        col_model, col_opts = st.columns([3, 2])
        with col_model:
            selected_display = st.selectbox(
                "Modello LLM", provider_options, index=selected_idx,
                key="model_selector_chat", label_visibility="collapsed",
            )
            idx = provider_options.index(selected_display)
            st.session_state.selected_provider = providers[idx][2]

        with col_opts:
            o1, o2 = st.columns(2)
            with o1:
                use_web = st.checkbox("Ricerca web", value=st.session_state.get("use_web_search", False), key="web_cb")
                st.session_state.use_web_search = use_web
            with o2:
                use_depth = st.checkbox("Approfondimento", value=st.session_state.get("use_depth_reasoning", False), key="depth_cb")
                st.session_state.use_depth_reasoning = use_depth

        # Storia chat
        _render_chat_history()

        # Domande suggerite (se disponibili)
        suggested = st.session_state.get("suggested_questions", [])
        if suggested:
            st.markdown("**💡 Domande suggerite:**")
            cols = st.columns(min(len(suggested), 3))
            for i, sq in enumerate(suggested):
                with cols[i % 3]:
                    if st.button(sq, key=f"suggest_{i}", use_container_width=True):
                        st.session_state.suggested_questions = []
                        st.session_state.chat_history.append({"role": "user", "content": sq})
                        st.session_state._pending_question = sq
                        st.rerun()

        # Input
        user_input = st.chat_input("Fai una domanda sui tuoi documenti...")

        # Gestisci domanda pendente (da suggerite) o nuovo input
        pending = st.session_state.pop("_pending_question", None)
        active_question = pending or user_input

        if active_question:
            if not pending:
                st.session_state.chat_history.append({"role": "user", "content": active_question})

            with st.chat_message("user"):
                st.write(active_question)

            with st.chat_message("assistant"):
                try:
                    docs = retrieve_relevant_docs(active_question)
                    st.session_state.last_retrieved_docs = docs

                    # Confidence score
                    confidence = calculate_confidence(docs, active_question)
                    st.session_state.last_confidence = confidence

                    if not docs:
                        st.warning("Nessun documento rilevante trovato.")
                        st.session_state.chat_history.append({
                            "role": "assistant",
                            "content": "Nessun documento rilevante trovato nell'indice.",
                        })
                    else:
                        current_role = get_user_role()
                        history_for_prompt = st.session_state.chat_history[:-1]

                        # Chiamata con streaming
                        result = rag_query(
                            active_question,
                            provider_override=st.session_state.get("selected_provider"),
                            use_web_search=st.session_state.get("use_web_search", False),
                            use_depth_reasoning=st.session_state.get("use_depth_reasoning", False),
                            user_role=current_role,
                            chat_history=history_for_prompt,
                            search_filters=st.session_state.get("search_filters"),
                            stream=True,
                        )

                        # rag_query restituisce (generator_or_answer, timing)
                        answer_gen, timing = result

                        # Streaming output
                        answer = st.write_stream(answer_gen)

                        if not answer or (isinstance(answer, str) and answer.strip() == ""):
                            answer = "Errore: risposta vuota dal modello."

                        # Timing info
                        st.session_state.last_timing = timing
                        total = timing.get("total", timing.get("pre_llm_total", "?"))
                        if timing.get("cache_hit"):
                            st.markdown(f"<span class='timing-badge'>⚡ Cache — {total}s</span>", unsafe_allow_html=True)
                        else:
                            st.markdown(f"<span class='timing-badge'>⏱ {total}s</span>", unsafe_allow_html=True)

                        # Citazioni
                        citations = parse_citations(answer)
                        citations = match_citations_to_sources(citations, docs)
                        st.session_state.last_citations = citations

                        st.session_state.chat_history.append({
                            "role": "assistant",
                            "content": answer,
                        })

                        # Genera domande suggerite (in background)
                        try:
                            suggestions = generate_suggested_questions(active_question, answer)
                            st.session_state.suggested_questions = suggestions
                        except Exception:
                            st.session_state.suggested_questions = []

                except Exception as e:
                    import traceback
                    st.error(f"Errore: {str(e)}")
                    with st.expander("Dettagli errore"):
                        st.code(traceback.format_exc())
                    st.session_state.chat_history.append({
                        "role": "assistant", "content": f"Errore: {str(e)}",
                    })

            st.rerun()

        # Azioni
        if st.session_state.chat_history:
            col_clear, col_export = st.columns(2)
            with col_clear:
                if st.button("🗑️ Pulisci Chat", key="clear_chat", use_container_width=True):
                    st.session_state.chat_history = []
                    st.session_state.last_retrieved_docs = []
                    st.session_state.last_citations = []
                    st.session_state.last_confidence = None
                    st.session_state.suggested_questions = []
                    st.session_state.last_timing = None
                    st.rerun()
            with col_export:
                word_data = _export_chat_as_word()
                if word_data:
                    st.download_button(
                        "📄 Esporta Word", data=word_data,
                        file_name="trillium_chat.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="export_word", use_container_width=True,
                    )
                else:
                    st.caption("Installa python-docx per export Word")
