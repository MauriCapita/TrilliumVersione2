import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from openai import OpenAI
import base64
import pandas as pd
import openpyxl
from docx import Document

# Supporto per HEIC/HEIF
try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
    HEIF_SUPPORT = True
except ImportError:
    HEIF_SUPPORT = False
from config import (
    PROVIDER,
    OPENAI_API_KEY,
    OPENROUTER_API_KEY,
    GOOGLE_CLOUD_VISION_KEY,
    GOOGLE_CLOUD_VISION_PROJECT,
    VISION_MODEL_OPENAI,
    VISION_MODEL_OPENROUTER,
    MIN_TEXT_LENGTH,
    MIN_TEXT_LENGTH_IMAGE,
    MAX_IMAGE_SIZE_MB,
    MAX_IMAGE_SIDE_PX,
    TESSERACT_LANG,
    IMAGE_EXTRACTION_STRATEGY,
)
from rich import print

# Statistiche estrazione immagini (per log/UI): lista di {"path", "method", "pages", "size_mb"}
IMAGE_EXTRACTION_STATS = []

# ============================================================
# CLIENT OPENAI / OPENROUTER
# ============================================================

def get_openai_client():
    return OpenAI(api_key=OPENAI_API_KEY)

def get_openrouter_client():
    return OpenAI(
        api_key=OPENROUTER_API_KEY,
        base_url="https://openrouter.ai/api/v1"
    )

# ============================================================
# UTILS
# ============================================================

def _resize_image_if_needed(img, max_side_px):
    """Ridimensiona PIL Image se il lato lungo supera max_side_px (mantiene aspect ratio)."""
    w, h = img.size
    if w <= max_side_px and h <= max_side_px:
        return img
    ratio = min(max_side_px / w, max_side_px / h)
    new_w, new_h = int(w * ratio), int(h * ratio)
    return img.resize((new_w, new_h), Image.Resampling.LANCZOS)


def _image_to_base64_for_vision(path, max_side_px=None, max_size_mb=None):
    """
    Carica immagine, ridimensiona se necessario (lato o dimensione file), restituisce (base64, mime_type).
    Se il file supera max_size_mb viene caricato con PIL, ridimensionato per stare sotto, poi codificato.
    """
    max_side_px = max_side_px or MAX_IMAGE_SIDE_PX
    max_size_mb = max_size_mb or MAX_IMAGE_SIZE_MB
    max_size_bytes = int(max_size_mb * 1024 * 1024)
    ext = os.path.splitext(path)[1].lower()
    file_size = os.path.getsize(path)

    try:
        img = Image.open(path)
        if img.mode not in ("RGB", "L", "1"):
            img = img.convert("RGB")
        img.load()
    except Exception as e:
        raise ValueError(f"Impossibile aprire immagine: {e}")

    w, h = img.size
    if w <= 0 or h <= 0:
        raise ValueError("Immagine con dimensioni nulle o non valide")
    if w > max_side_px or h > max_side_px:
        img = _resize_image_if_needed(img, max_side_px)

    if file_size > max_size_bytes:
        img = _resize_image_if_needed(img, 2048)

    buf = io.BytesIO()
    out_fmt = "PNG" if ext in [".tif", ".tiff", ".bmp", ".heic", ".heif"] else "JPEG"
    img.save(buf, format=out_fmt, quality=85)
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    mime = "image/png" if out_fmt == "PNG" else "image/jpeg"
    return b64, mime


def encode_image_to_base64(path):
    """Converte un file immagine in base64. Usa limite configurabile MAX_IMAGE_SIZE_MB."""
    max_size_bytes = int(MAX_IMAGE_SIZE_MB * 1024 * 1024)
    file_size = os.path.getsize(path)
    file_size_mb = file_size / (1024 * 1024)
    if file_size > max_size_bytes:
        raise ValueError(f"File troppo grande: {file_size_mb:.2f} MB > {MAX_IMAGE_SIZE_MB} MB. Impossibile elaborare.")
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")

# ============================================================
# 1) ESTRATTORI LOCALI
# ============================================================

def extract_pdf_local(path):
    """Estrae testo da PDF tramite PyMuPDF"""
    try:
        doc = fitz.open(path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text.strip()
    except Exception as e:
        print(f"[yellow]Errore PyMuPDF su PDF {path}: {e}[/yellow]")
        return ""

def _ocr_single_image(img, lang=None):
    """Esegue OCR su una singola PIL Image; inversione per WhiteIsZero (tag TIFF 262 = 0)."""
    lang = lang or TESSERACT_LANG
    tags = getattr(img, "tag_v2", {}) or {}
    if img.mode in ("1", "L") and tags.get(262) == 0:
        try:
            from PIL import ImageOps
            img = ImageOps.invert(img)
        except Exception:
            pass
    if img.mode not in ("RGB", "RGBA", "L", "1"):
        img = img.convert("RGB")
    try:
        return pytesseract.image_to_string(img, lang=lang).strip()
    except Exception:
        return pytesseract.image_to_string(img).strip()


def extract_tif_local(path):
    """OCR locale con Tesseract; supporta TIF multi-pagina e bi-level (WhiteIsZero)."""
    try:
        img = Image.open(path)
        n_frames = getattr(img, "n_frames", 1)
        parts = []
        for i in range(n_frames):
            try:
                img.seek(i)
            except EOFError:
                break
            page_text = _ocr_single_image(img)
            if page_text:
                parts.append(f"--- Page {i + 1} ---\n{page_text}")
        return "\n\n".join(parts) if parts else ""
    except Exception as e:
        print(f"[yellow]Errore Tesseract su TIF {path}: {e}[/yellow]")
        return ""


def extract_image_local(path):
    """OCR locale per immagini (BMP, PNG, HEIC, HEIF) con Tesseract; lingua configurabile."""
    try:
        ext = os.path.splitext(path)[1].lower()
        if ext in [".heic", ".heif"]:
            if not HEIF_SUPPORT:
                print(f"[yellow]⚠ Supporto HEIC/HEIF non disponibile. Installa: pip install pillow-heif[/yellow]")
                return ""
        img = Image.open(path)
        if img.mode != "RGB":
            img = img.convert("RGB")
        text = pytesseract.image_to_string(img, lang=TESSERACT_LANG)
        return text.strip()
    except Exception as e:
        print(f"[yellow]Errore OCR su immagine {path}: {e}[/yellow]")
        return ""

def extract_log_local(path):
    """Estrae testo da file .log (file di testo semplice)"""
    try:
        # Prova con encoding UTF-8
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # Se UTF-8 fallisce, prova con latin-1 o altri encoding comuni
            try:
                with open(path, 'r', encoding='latin-1') as f:
                    text = f.read()
            except UnicodeDecodeError:
                # Ultimo tentativo con error handling
                with open(path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
        
        return text.strip()
    except Exception as e:
        print(f"[yellow]Errore lettura file .log {path}: {e}[/yellow]")
        return ""

def extract_txt_local(path):
    """Estrae testo da file .txt (file di testo semplice)"""
    try:
        # Prova con encoding UTF-8
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # Se UTF-8 fallisce, prova con latin-1 o altri encoding comuni
            try:
                with open(path, 'r', encoding='latin-1') as f:
                    text = f.read()
            except UnicodeDecodeError:
                # Ultimo tentativo con error handling
                with open(path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
        
        return text.strip()
    except Exception as e:
        print(f"[yellow]Errore lettura file .txt {path}: {e}[/yellow]")
        return ""

def extract_excel_local(path):
    """Estrae testo da file Excel (.xlsx, .xls) mantenendo la struttura, formule, commenti e nomi"""
    try:
        text_parts = []
        
        # Usa openpyxl per accedere a formule, commenti e nomi
        try:
            wb = openpyxl.load_workbook(path, data_only=False)  # data_only=False per vedere formule
            
            # Estrai nomi definiti (named ranges)
            if wb.defined_names:
                text_parts.append("=== DEFINED NAMES (Named Ranges) ===\n")
                for name, definition in wb.defined_names.items():
                    text_parts.append(f"Name: {name} = {definition}")
                text_parts.append("\n")
            
            # Leggi tutte le sheet
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                text_parts.append(f"\n=== SHEET: {sheet_name} ===\n")
                
                # Estrai commenti
                comments_found = []
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.comment:
                            comments_found.append(f"Cell {cell.coordinate}: {cell.comment.text}")
                
                if comments_found:
                    text_parts.append("--- COMMENTS ---\n")
                    text_parts.extend(comments_found)
                    text_parts.append("\n")
                
                # Estrai dati con formule quando disponibili
                # Prima prova con pandas per struttura tabellare
                try:
                    df = pd.read_excel(path, sheet_name=sheet_name, header=0, engine='openpyxl')
                    
                    # Se ci sono colonne con nomi, usa quelli
                    if not df.columns.empty and any(pd.notna(df.columns)):
                        # Header row
                        header_row = " | ".join([str(col) if pd.notna(col) else "" for col in df.columns])
                        text_parts.append("--- DATA TABLE ---\n")
                        text_parts.append(header_row)
                        text_parts.append("-" * min(len(header_row), 100))
                        
                        # Righe dati
                        for idx, row in df.iterrows():
                            row_values = []
                            for col in df.columns:
                                cell_value = row[col]
                                if pd.notna(cell_value):
                                    row_values.append(str(cell_value))
                                else:
                                    row_values.append("")
                            
                            if any(row_values):
                                text_parts.append(" | ".join(row_values))
                    
                    text_parts.append("\n")
                    
                except:
                    pass
                
                # Estrai formule importanti (celle con formule, non solo valori)
                formulas_found = []
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.data_type == 'f' and cell.value:  # Formula
                            # Prendi anche il valore calcolato se disponibile
                            formula = str(cell.value)
                            try:
                                calculated_value = cell.value if cell.data_type != 'f' else ws[cell.coordinate].value
                                if calculated_value is not None:
                                    formulas_found.append(f"Cell {cell.coordinate}: {formula} = {calculated_value}")
                                else:
                                    formulas_found.append(f"Cell {cell.coordinate}: {formula}")
                            except:
                                formulas_found.append(f"Cell {cell.coordinate}: {formula}")
                
                if formulas_found:
                    text_parts.append("--- FORMULAS ---\n")
                    text_parts.extend(formulas_found[:100])  # Limita a 100 formule per non esagerare
                    text_parts.append("\n")
                
                # Estrai anche tutte le celle con valori significativi (fallback completo)
                significant_cells = []
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value is not None and str(cell.value).strip():
                            # Se non è già stato incluso nelle tabelle
                            cell_info = f"Cell {cell.coordinate}: {cell.value}"
                            if cell.data_type == 'f':
                                cell_info += f" [FORMULA: {cell.value}]"
                            significant_cells.append(cell_info)
                
                # Aggiungi solo se ci sono celle significative non già incluse
                if significant_cells and len(significant_cells) < 500:  # Limita per file molto grandi
                    text_parts.append("--- ALL SIGNIFICANT CELLS ---\n")
                    text_parts.extend(significant_cells[:200])  # Limita a 200 celle
                    text_parts.append("\n")
            
            wb.close()
            
        except Exception as e:
            # Fallback: usa solo pandas se openpyxl fallisce
            excel_file = pd.ExcelFile(path)
            
            for sheet_name in excel_file.sheet_names:
                text_parts.append(f"\n=== SHEET: {sheet_name} ===\n")
                
                try:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name, header=0)
                    
                    if not df.columns.empty and any(pd.notna(df.columns)):
                        header_row = " | ".join([str(col) if pd.notna(col) else "" for col in df.columns])
                        text_parts.append(header_row)
                        text_parts.append("-" * min(len(header_row), 100))
                    
                    for idx, row in df.iterrows():
                        row_values = []
                        for col in df.columns:
                            cell_value = row[col]
                            if pd.notna(cell_value):
                                row_values.append(str(cell_value))
                            else:
                                row_values.append("")
                        
                        if any(row_values):
                            text_parts.append(" | ".join(row_values))
                    
                except:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
                    
                    for idx, row in df.iterrows():
                        row_text = []
                        for col_idx, cell_value in enumerate(row):
                            if pd.notna(cell_value):
                                row_text.append(str(cell_value))
                        
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                text_parts.append("\n")
        
        # Unisci tutto il testo
        full_text = "\n".join(text_parts)
        
        # Rimuovi righe vuote eccessive ma mantieni struttura
        lines = []
        prev_empty = False
        for line in full_text.split("\n"):
            if line.strip():
                lines.append(line)
                prev_empty = False
            elif not prev_empty:
                lines.append("")
                prev_empty = True
        
        return "\n".join(lines).strip()
        
    except ImportError:
        print(f"[yellow]⚠ Librerie Excel non installate. Installa con: pip install openpyxl pandas[/yellow]")
        return ""
    except Exception as e:
        print(f"[yellow]Errore estrazione Excel su {path}: {e}[/yellow]")
        return ""

def extract_word_local(path):
    """Estrae testo da documenti Word (.docx, .doc) mantenendo la struttura"""
    try:
        text_parts = []
        
        # Apri il documento Word
        doc = Document(path)
        
        # Estrai tutti i paragrafi
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Estrai testo dalle tabelle
        for table in doc.tables:
            text_parts.append("\n=== TABELLA ===\n")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
            text_parts.append("\n")
        
        # Unisci tutto il testo
        full_text = "\n".join(text_parts)
        
        # Rimuovi righe vuote eccessive
        lines = []
        prev_empty = False
        for line in full_text.split("\n"):
            if line.strip():
                lines.append(line)
                prev_empty = False
            elif not prev_empty:
                lines.append("")
                prev_empty = True
        
        return "\n".join(lines).strip()
        
    except ImportError:
        print(f"[yellow]⚠ Libreria python-docx non installata. Installa con: pip install python-docx[/yellow]")
        return ""
    except Exception as e:
        # Per file .doc vecchi (non .docx), potrebbe servire un'altra libreria
        if path.lower().endswith('.doc') and not path.lower().endswith('.docx'):
            print(f"[yellow]⚠ File .doc vecchio non supportato direttamente. Converti in .docx o usa un convertitore.[/yellow]")
        print(f"[yellow]Errore estrazione Word su {path}: {e}[/yellow]")
        return ""

# ============================================================
# 2) ESTRATTORI GPT VISION (OPENAI / OPENROUTER)
# ============================================================

def extract_with_openai_vision(path):
    if not OPENAI_API_KEY:
        print("[yellow]⚠ OpenAI API key non configurata. Salto OpenAI Vision.[/yellow]")
        return ""
    try:
        client = get_openai_client()
        ext = os.path.splitext(path)[1].lower()
        if ext in [".tif", ".tiff", ".bmp", ".heic", ".heif"]:
            print(f"[yellow]⚠ OpenAI Vision non supporta {ext.upper()}. Usa un altro servizio.[/yellow]")
            return ""
        b64, mime_type = _image_to_base64_for_vision(path)
        response = client.chat.completions.create(
            model=VISION_MODEL_OPENAI,
            messages=[
                {"role": "user",
                 "content": [
                     {"type": "text", "text": "Analyze this technical engineering drawing and provide a comprehensive description that includes:\n\n- All text, labels, and annotations visible in the image\n- All numbers, dimensions, measurements, and tolerances\n- All notes, specifications, and technical information\n- Material information, part numbers, revision numbers\n- All symbols, codes, and references\n\nBe extremely thorough and include every detail you can see. This description will be used for document search and retrieval."},
                     {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{b64}"}}
                 ]}
            ],
            temperature=0.1,
            max_tokens=16000
        )
        return response.choices[0].message.content.strip()
    except ValueError as e:
        print(f"[red]✗ {str(e)}[/red]")
        return ""
    except Exception as e:
        print(f"[red]Errore OpenAI Vision su {path}: {e}[/red]")
        return ""

def extract_with_openrouter_vision(path, model=None):
    """Estrae testo usando OpenRouter; immagine ridimensionata se supera MAX_IMAGE_SIDE_PX o MAX_IMAGE_SIZE_MB."""
    if not OPENROUTER_API_KEY:
        print("[yellow]⚠ OpenRouter API key non configurata. Salto OpenRouter Vision.[/yellow]")
        return ""
    try:
        client = get_openrouter_client()
        b64, mime_type = _image_to_base64_for_vision(path)
        model_to_use = model or VISION_MODEL_OPENROUTER
        response = client.chat.completions.create(
            model=model_to_use,
            messages=[
                {"role": "user",
                 "content": [
                     {"type": "text", "text": "You are an OCR system. Extract ALL text visible in this image/document. Include every word, number, symbol, and character you can see. Preserve the structure and formatting as much as possible. Output ONLY the extracted text, nothing else."},
                     {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{b64}"}}
                 ]}
            ],
            temperature=0
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"[red]Errore OpenRouter Vision ({model or VISION_MODEL_OPENROUTER}) su {path}: {e}[/red]")
        return ""

def extract_with_google_vision(path):
    """Estrae testo usando Google Cloud Vision API (supporta TIF nativamente)"""
    # Controlla se le credenziali sono configurate
    if not GOOGLE_CLOUD_VISION_KEY and not GOOGLE_CLOUD_VISION_PROJECT:
        # Controlla anche la variabile d'ambiente standard
        if not os.getenv("GOOGLE_APPLICATION_CREDENTIALS"):
            print("[yellow]⚠ Google Cloud Vision non configurato. Salto Google Vision.[/yellow]")
            return ""
    
    try:
        from google.cloud import vision
        
        # Inizializza il client Google Vision
        # Google Cloud Vision richiede un file JSON delle credenziali o variabili d'ambiente
        if GOOGLE_CLOUD_VISION_KEY and os.path.exists(GOOGLE_CLOUD_VISION_KEY):
            # Se GOOGLE_CLOUD_VISION_KEY è un percorso a file JSON
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = GOOGLE_CLOUD_VISION_KEY
            client = vision.ImageAnnotatorClient()
        elif GOOGLE_CLOUD_VISION_PROJECT:
            # Usa le credenziali di default se il progetto è configurato
            # Assumendo che GOOGLE_APPLICATION_CREDENTIALS sia già impostato
            client = vision.ImageAnnotatorClient()
        else:
            # Prova con credenziali di default (da variabile d'ambiente)
            client = vision.ImageAnnotatorClient()
        
        # Usa immagine eventualmente ridimensionata (rispetta MAX_IMAGE_SIDE_PX e MAX_IMAGE_SIZE_MB)
        b64, _ = _image_to_base64_for_vision(path)
        content = base64.b64decode(b64)
        image = vision.Image(content=content)
        
        # Esegui OCR
        response = client.text_detection(image=image)
        texts = response.text_annotations
        
        if texts:
            # Il primo elemento contiene tutto il testo
            return texts[0].description.strip()
        else:
            return ""
            
    except ImportError:
        print(f"[yellow]⚠ google-cloud-vision non installato. Installa con: pip install google-cloud-vision[/yellow]")
        return ""
    except Exception as e:
        print(f"[red]Errore Google Vision su {path}: {e}[/red]")
        return ""

# ============================================================
# 3) PIPELINE IBRIDA COMPLETA
# ============================================================

def _fallback_description_for_image(path):
    """Testo minimo per ricerca: nome file + eventuale breve descrizione Vision."""
    name = os.path.basename(path)
    ext = os.path.splitext(path)[1].lower()
    base = f"Document: {name}. Extension: {ext}. Technical drawing or image."
    if OPENROUTER_API_KEY:
        try:
            b64, mime = _image_to_base64_for_vision(path)
            client = get_openrouter_client()
            r = client.chat.completions.create(
                model="anthropic/claude-3.5-sonnet",
                messages=[
                    {"role": "user", "content": [
                        {"type": "text", "text": "Describe briefly this image for document search: part number if visible, document type, main content in one sentence. Output only this short description, nothing else."},
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}
                    ]}
                ],
                temperature=0
            )
            desc = (r.choices[0].message.content or "").strip()
            if desc:
                base += " " + desc
        except Exception:
            pass
    return base


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    is_image = ext in [".tif", ".tiff", ".bmp", ".png", ".heic", ".heif", ".jpg", ".jpeg"]
    min_len = MIN_TEXT_LENGTH_IMAGE if is_image else MIN_TEXT_LENGTH
    file_size_mb = os.path.getsize(path) / (1024 * 1024)

    print(f"[cyan]➜ Estrazione testo da: {path}[/cyan]")

    # --- 1) TENTATIVO LOCALE (salta se vision_only per immagini) ---
    text_local = ""
    if not (is_image and IMAGE_EXTRACTION_STRATEGY == "vision_only"):
        if ext == ".pdf":
            text_local = extract_pdf_local(path)
        elif ext in [".tif", ".tiff"]:
            text_local = extract_tif_local(path)
        elif ext in [".xlsx", ".xls", ".xlsm"]:
            text_local = extract_excel_local(path)
        elif ext in [".docx", ".doc"]:
            text_local = extract_word_local(path)
        elif ext in [".bmp", ".png", ".heic", ".heif"]:
            text_local = extract_image_local(path)
        elif ext == ".log":
            text_local = extract_log_local(path)
        elif ext == ".txt":
            text_local = extract_txt_local(path)

    if len(text_local) >= min_len:
        print("[green]✔ Testo estratto localmente[/green]")
        if is_image:
            n_pages = 1
            if ext in [".tif", ".tiff"]:
                try:
                    with Image.open(path) as im:
                        n_pages = getattr(im, "n_frames", 1)
                except Exception:
                    pass
            IMAGE_EXTRACTION_STATS.append({"path": path, "method": "local", "pages": n_pages, "size_mb": round(file_size_mb, 2)})
        return text_local

    if ext == ".txt":
        print(f"[yellow]⚠ File .txt con testo insufficiente. Salto Vision.[/yellow]")
        return text_local

    if is_image and IMAGE_EXTRACTION_STRATEGY == "local_only":
        text_fallback = _fallback_description_for_image(path)
        IMAGE_EXTRACTION_STATS.append({"path": path, "method": "fallback", "pages": 1, "size_mb": round(file_size_mb, 2)})
        return text_fallback

    print("[yellow]⚠ Testo insufficiente localmente → passo a Vision[/yellow]")

    # --- 2) PER IMMAGINI: Vision (Google, Claude, Gemini) ---
    result_text = ""
    method_used = None
    if ext in [".tif", ".tiff", ".bmp", ".png", ".heic", ".heif"]:
        if GOOGLE_CLOUD_VISION_KEY or GOOGLE_CLOUD_VISION_PROJECT:
            print(f"[cyan]→ Tentativo con Google Cloud Vision API...[/cyan]")
            result_text = extract_with_google_vision(path)
            if len(result_text) >= min_len:
                method_used = "vision_google"
                print("[green]✔ Testo estratto con Google Vision API[/green]")
        if not result_text or len(result_text) < min_len:
            if OPENROUTER_API_KEY:
                print(f"[cyan]→ Tentativo con Claude 3.5 Sonnet (OpenRouter)...[/cyan]")
                result_text = extract_with_openrouter_vision(path, model="anthropic/claude-3.5-sonnet")
                if len(result_text) >= min_len:
                    method_used = "vision_claude"
                    print("[green]✔ Testo estratto con Claude 3.5 Sonnet[/green]")
            if not result_text or len(result_text) < min_len:
                if OPENROUTER_API_KEY:
                    print(f"[cyan]→ Tentativo con Gemini 2.0 Flash (OpenRouter)...[/cyan]")
                    result_text = extract_with_openrouter_vision(path, model="google/gemini-2.0-flash-exp")
                    if len(result_text) >= min_len:
                        method_used = "vision_gemini"
                        print("[green]✔ Testo estratto con Gemini[/green]")
        if result_text and len(result_text) >= min_len:
            n_pages = 1
            try:
                with Image.open(path) as im:
                    n_pages = getattr(im, "n_frames", 1)
            except Exception:
                pass
            IMAGE_EXTRACTION_STATS.append({"path": path, "method": method_used or "vision", "pages": n_pages, "size_mb": round(file_size_mb, 2)})
            return result_text
        if not result_text:
            print(f"[yellow]⚠ Nessun servizio cloud disponibile per {ext.upper()}.[/yellow]")

    # --- 3) ALTRI FORMATI: OpenAI Vision ---
    if not result_text or len(result_text) < min_len:
        text_ai = extract_with_openai_vision(path)
        if len(text_ai) >= min_len:
            if is_image:
                IMAGE_EXTRACTION_STATS.append({"path": path, "method": "vision_openai", "pages": 1, "size_mb": round(file_size_mb, 2)})
            return text_ai
        if OPENROUTER_API_KEY:
            text_or = extract_with_openrouter_vision(path)
            if len(text_or) >= min_len:
                if is_image:
                    IMAGE_EXTRACTION_STATS.append({"path": path, "method": "vision_openrouter", "pages": 1, "size_mb": round(file_size_mb, 2)})
                return text_or

    # --- 4) FALLBACK per immagini: descrizione minima ---
    if is_image:
        result_text = (result_text or text_local or "").strip()
        if len(result_text) < min_len:
            result_text = _fallback_description_for_image(path)
            print("[green]✔ Usata descrizione di fallback per ricerca[/green]")
        n_pages = 1
        try:
            with Image.open(path) as im:
                n_pages = getattr(im, "n_frames", 1)
        except Exception:
            pass
        IMAGE_EXTRACTION_STATS.append({"path": path, "method": "fallback", "pages": n_pages, "size_mb": round(file_size_mb, 2)})
        return result_text

    print("[red]✘ Nessun sistema è riuscito a estrarre testo[/red]")
    return ""


def get_image_extraction_stats():
    """Restituisce e azzera la lista delle statistiche di estrazione immagini (per log/UI)."""
    global IMAGE_EXTRACTION_STATS
    out = list(IMAGE_EXTRACTION_STATS)
    IMAGE_EXTRACTION_STATS = []
    return out


def clear_image_extraction_stats():
    """Azzera le statistiche di estrazione immagini."""
    global IMAGE_EXTRACTION_STATS
    IMAGE_EXTRACTION_STATS = []

